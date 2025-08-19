import os
import aiofiles
from typing import Union, Dict, Any
import PyPDF2
import docx
from io import BytesIO
from fastapi import UploadFile
from ..core.config import settings

class DocumentProcessor:
    def __init__(self):
        self.supported_extensions = {'.pdf', '.docx', '.doc', '.txt'}
        os.makedirs(settings.upload_dir, exist_ok=True)
    
    async def process_uploaded_file(self, file: UploadFile, domain: str) -> Dict[str, Any]:
        """Process uploaded file and extract text content"""
        
        # Validate file
        self._validate_file(file)
        
        # Save file temporarily
        file_path = await self._save_uploaded_file(file)
        
        try:
            # Extract content based on file type
            content = await self._extract_content(file_path, file.filename)
            
            return {
                'filename': file.filename,
                'content': content,
                'domain': domain,
                'file_size': os.path.getsize(file_path)
            }
        
        finally:
            # Clean up temporary file
            if os.path.exists(file_path):
                os.remove(file_path)
    
    def _validate_file(self, file: UploadFile) -> None:
        """Validate uploaded file"""
        
        # Check file extension
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_ext}. Supported types: {self.supported_extensions}")
        
        # More flexible content type validation
        if file.content_type is None:
            # If content type is None, rely on file extension
            print(f"Warning: Content type is None for file {file.filename}, using extension validation")
            return
        
        # Allow common PDF content types
        allowed_content_types = {
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword',
            'text/plain',
            'application/octet-stream',  # Add this for binary files
            'binary/octet-stream'        # Add this for binary files
        }
        
        if file.content_type not in allowed_content_types:
            # Log the actual content type for debugging
            print(f"Warning: Unexpected content type: {file.content_type} for file {file.filename}")
            # Don't fail validation, just log a warning
            return
    
    async def _save_uploaded_file(self, file: UploadFile) -> str:
        """Save uploaded file temporarily"""
        
        import uuid
        file_id = str(uuid.uuid4())
        file_ext = os.path.splitext(file.filename)[1]
        file_path = os.path.join(settings.upload_dir, f"{file_id}{file_ext}")
        
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            if len(content) > settings.max_file_size:
                raise ValueError(f"File too large. Maximum size: {settings.max_file_size} bytes")
            await f.write(content)
        
        return file_path
    
    async def _extract_content(self, file_path: str, filename: str) -> str:
        """Extract text content from file"""
        
        file_ext = os.path.splitext(filename)[1].lower()
        
        if file_ext == '.pdf':
            return await self._extract_pdf_content(file_path)
        elif file_ext in ['.docx', '.doc']:
            return await self._extract_docx_content(file_path)
        elif file_ext == '.txt':
            return await self._extract_txt_content(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {file_ext}")
    
    async def _extract_pdf_content(self, file_path: str) -> str:
        """Extract text from PDF file"""
        
        try:
            content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            content += f"\n--- Page {page_num + 1} ---\n"
                            content += page_text
                    except Exception as e:
                        print(f"Warning: Could not extract text from page {page_num + 1}: {e}")
                        continue
            
            if not content.strip():
                raise ValueError("No text content could be extracted from the PDF")
            
            return content.strip()
        
        except Exception as e:
            raise ValueError(f"Failed to extract PDF content: {str(e)}")
    
    async def _extract_docx_content(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        
        try:
            doc = docx.Document(file_path)
            content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content += " | ".join(row_text) + "\n"
            
            if not content.strip():
                raise ValueError("No text content could be extracted from the document")
            
            return content.strip()
        
        except Exception as e:
            raise ValueError(f"Failed to extract DOCX content: {str(e)}")
    
    async def _extract_txt_content(self, file_path: str) -> str:
        """Extract text from TXT file"""
        
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()
            
            if not content.strip():
                raise ValueError("The text file is empty")
            
            return content.strip()
        
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                async with aiofiles.open(file_path, 'r', encoding='latin-1') as file:
                    content = await file.read()
                return content.strip()
            except Exception as e:
                raise ValueError(f"Failed to read text file with UTF-8 or Latin-1 encoding: {str(e)}")
        
        except Exception as e:
            raise ValueError(f"Failed to extract text content: {str(e)}")
    
    def get_file_info(self, file: UploadFile) -> Dict[str, Any]:
        """Get basic information about uploaded file"""
        
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        return {
            'filename': file.filename,
            'content_type': file.content_type,
            'file_extension': file_ext,
            'is_supported': file_ext in self.supported_extensions
        } 